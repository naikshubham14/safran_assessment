import streamlit as st
from rule_checker import RuleChecker
from dotenv import load_dotenv
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import base64
import spacy
from spacy.cli import download

# Load environment variables (for Gemini API key)
load_dotenv()

# Page configuration
st.set_page_config(page_title="Rule Checker", page_icon="🔍", layout="wide")

def main()
    def write_to_docx(results: dict, output_file="output.docx"):
        """
        Takes the results of the rule checker and writes them to a docx file
        Args:
            dict: The rule check results
            
        Returns:
            str, int: path to the output file, total number of violations
        """
        # Create a new Word document
        doc = Document()
        
        # Initialize variables
        total_violations = 0

        # Process results
        for result in results:
            paragraph = doc.add_paragraph(result["sentence"])
            annotation = []
            
            # Check for violations and prepare superscript
            for i, v in enumerate(result["violations"]):
                if v:
                    total_violations += 1
                    annotation.append(str(i + 1))
            
            # Add superscript annotations if violations exist
            if annotation:
                superscript_text = f"{','.join(annotation)}"
                run = paragraph.add_run()
                add_superscript(run, superscript_text)
        
        # Save the document
        doc.save(output_file)
        
        # Return file path and total violations
        return output_file, total_violations

    # Helper function to add superscript text
    def add_superscript(run, text: str):
        """
        Adds superscript text to a run in a Word document.
        Args:
            run: The run object to which the superscript text will be added.
            text (str): The text to be added as superscript.
        """
        r = run._r
        for char in text:
            elem = OxmlElement("w:r")
            t_elem = OxmlElement("w:t")
            t_elem.text = char
            elem.append(t_elem)
            r.append(elem)
            if char.isdigit():
                rPr = elem.get_or_add_rPr()
                vert_align = OxmlElement("w:vertAlign")
                vert_align.set(qn("w:val"), "superscript")
                rPr.append(vert_align)

    def create_download_link(file_path, button_text="Download Word File"):
        """
        Creates a download link for a file.
        
        Args:
            file_path (str): The path to the file to be downloaded.
            button_text (str): The text to display on the download button.
            
        Returns:
            str: An HTML anchor tag that allows the user to download the file.
        """
        
        with open(file_path, "rb") as f:
            file_data = f.read()
        encoded_file = base64.b64encode(file_data).decode()
        href = f'<a href="data:application/octet-stream;base64,{encoded_file}" download="{file_path}">{button_text}</a>'
        return href


    # Title and description
    st.title("Rule Checker Application")
    st.markdown("""
        This application checks if your text adheres to the following rules:
        1. Use articles/demonstratives before nouns.
        2. Use active voice in procedural writing.
        3. Write one instruction per sentence (unless simultaneous).
        4. Use imperative (command) form.
        5. Keep sentences under 20 words.
    """)

    # Initialize RuleChecker
    rule_checker = RuleChecker(gemini_api_key=os.getenv("GEMINI_API_KEY"))

    # Input options
    st.sidebar.header("Input Options")
    input_type = st.sidebar.radio("Choose input type:", ("Text", "File"))

    # Text input
    # Text input
    if input_type == "Text":
        text = st.text_area("Enter your text here:", height=200)
        if st.button("Check Rules"):
            if text.strip():
                results = rule_checker.check_rules(text)
                output = ""
                total_violations = 0
                for result in results:
                    output += result["sentence"]
                    annotation = []
                    for i,v in enumerate(result["violations"]):
                        if v:
                            total_violations += 1
                            annotation.append(str(i+1))             
                    if len(annotation) > 0:
                        output += f"<sup>{",".join(annotation)}</sup> "
                    else:
                        output += " " 
                if total_violations == 0:  
                    st.success("No violations.")
                else:      
                    st.subheader("Output")
                    st.markdown(output, unsafe_allow_html=True)

    # File input
    else:
        uploaded_file = st.file_uploader("Upload a text file", type=["txt", "pdf"])
        if uploaded_file is not None:
            text = uploaded_file.read().decode("utf-8")
            st.text_area("File content:", text, height=200)
            if st.button("Check Rules"):
                if text.strip():
                    results = rule_checker.check_rules(text)
                    output_file, total_violations = write_to_docx(results)
                    if total_violations == 0:
                        st.success("No violations.")
                    else:
                        st.markdown(create_download_link(output_file), unsafe_allow_html=True)

if __name__ == "__main__":
    
    model_name = "en_core_web_sm"
    try:
        # Attempt to load the model
        spacy.load(model_name)
        print(f"Model '{model_name}' is already downloaded.")
    except OSError:
        # If the model is not found, download it
        print(f"Model '{model_name}' not found. Downloading...")
        download(model_name)
        print(f"Model '{model_name}' has been downloaded.")
    main()
